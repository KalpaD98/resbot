import csv
import json
import os

from docx import Document

# Name of the output CSV file
csv_file_name = "output.csv"

# Define directory names
directories = [
    "config-00-default",
    "config-01",
    "config-02-gpt",
    "config-03-gpt-2",
    "config-04-bert",
    "config-05-roberta",
    "config-06-distillbert",
    "config-07-xlnet",
    "config-08-spacy-lg",
    "config-09-spacy-trf",
]

# List to store the extracted values
data = []

# Function to extract scores from JSON file
def get_scores_from_json(json_file_path):
    with open(json_file_path, "r") as json_file:
        json_data = json.load(json_file)
        weighted_avg = json_data["weighted avg"]
        precision = round(weighted_avg["precision"] * 100, 2)
        recall = round(weighted_avg["recall"] * 100, 2)
        f1_score = round(weighted_avg["f1-score"] * 100, 2)
    return precision, recall, f1_score

# Loop over the directories
for directory in directories:
    # Normal config directory
    json_file_path = os.path.join(directory, "intent_report.json")
    precision, recall, f1_score = get_scores_from_json(json_file_path)

    # 'With-zero-shot' config directory
    zero_shot_dir = f"{directory}-with-zero-shot"
    json_file_path_zero_shot = os.path.join(zero_shot_dir, "intent_report.json")
    precision_zero_shot, recall_zero_shot, f1_score_zero_shot = get_scores_from_json(json_file_path_zero_shot)

    # Calculate the difference between f1 scores
    f1_score_difference = round(f1_score_zero_shot - f1_score, 2)

    # Append to data
    data.append([
        directory, precision, recall, f1_score,
        zero_shot_dir, precision_zero_shot, recall_zero_shot, f1_score_zero_shot, f1_score_difference
    ])

# Sort data by f1-score of Zero Shot Directory in descending order
data.sort(key=lambda x: x[7], reverse=True)

# Format data as string with percentages
data = [[row[0], f"{row[1]}%", f"{row[2]}%", f"{row[3]}%", row[4], f"{row[5]}%", f"{row[6]}%", f"{row[7]}%", f"{row[8]}%"] for row in data]

# Write the data to the CSV file
with open(csv_file_name, "w", newline="") as csv_file:
    writer = csv.writer(csv_file)
    writer.writerow(["Config Directory", "Precision", "Recall", "F1-Score", "Zero Shot Directory", "Zero Shot Precision", "Zero Shot Recall", "Zero Shot F1-Score", "F1-Score Difference"])  # Write the header
    writer.writerows(data)  # Write the data

# Start a new document
doc = Document()

# Add a title
doc.add_heading('Config Comparison Results', 0)

# Define the table structure
table = doc.add_table(rows=1, cols=8)

# Define the headers
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Config Name'
hdr_cells[1].text = 'Config Precision'
hdr_cells[2].text = 'Config Recall'
hdr_cells[3].text = 'Config F1-Score'
hdr_cells[4].text = 'Hybrid Config Precision'
hdr_cells[5].text = 'Hybrid Config Recall'
hdr_cells[6].text = 'Hybrid Config F1-Score'
hdr_cells[7].text = 'F1-Score Difference'

# Fill the table with data from your CSV
with open(csv_file_name, 'r') as f:
    csv_reader = csv.reader(f)
    next(csv_reader, None)  # Skip the header
    for row in csv_reader:
        cells = table.add_row().cells
        config_name = row[0]
        cells[0].text = config_name
        cells[1].text = row[1]
        cells[2].text = row[2]
        cells[3].text = row[3]
        cells[4].text = row[5]
        cells[5].text = row[6]
        cells[6].text = row[7]
        cells[7].text = row[8]

# Save the document
doc.save('config_results.docx')
